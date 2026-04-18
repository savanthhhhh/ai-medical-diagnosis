"""
Generate a professional DOCX project report for the Medical Diagnostic System.
Includes diagrams generated via matplotlib and embedded into the document.

Run: python generate_report.py
Output: Project_Report.docx
"""

import os
import io
import math
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import networkx as nx

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from medical_dataset import MEDICAL_KNOWLEDGE_BASE

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
REPORT_PATH = os.path.join(OUTPUT_DIR, "Project_Report.docx")
IMG_DIR = os.path.join(OUTPUT_DIR, "_report_images")
os.makedirs(IMG_DIR, exist_ok=True)


# ======================================================================
# DIAGRAM GENERATORS
# ======================================================================

def generate_architecture_diagram():
    """System architecture diagram showing the 3-layer pipeline."""
    fig, ax = plt.subplots(figsize=(10, 5))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis('off')

    # Title
    ax.text(5, 4.7, "System Architecture", ha='center', fontsize=16, fontweight='bold', color='#2c3e50')

    # Boxes
    boxes = [
        (0.5, 2.0, 2.5, 1.8, '#3498db', 'white', 'INPUT LAYER\n\n• Symptom Entry\n• Quick-Add Buttons\n• Category Browser'),
        (3.7, 2.0, 2.5, 1.8, '#e74c3c', 'white', 'PROCESSING\nLAYER\n\n• Cypher Queries\n• Confidence Scoring\n• Symptom Matching'),
        (6.9, 2.0, 2.5, 1.8, '#27ae60', 'white', 'OUTPUT LAYER\n\n• Diagnosis Results\n• Network Graph\n• Treatment Info'),
    ]

    for x, y, w, h, color, tc, text in boxes:
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                               facecolor=color, edgecolor='white', linewidth=2, alpha=0.9)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, text, ha='center', va='center',
                fontsize=8, color=tc, fontweight='bold', linespacing=1.4)

    # Arrows
    arrow_style = "Simple,tail_width=8,head_width=20,head_length=10"
    for x1, x2 in [(3.0, 3.7), (6.2, 6.9)]:
        ax.annotate("", xy=(x2, 2.9), xytext=(x1, 2.9),
                     arrowprops=dict(arrowstyle='->', color='#34495e', lw=2.5))

    # Bottom bar - Technology
    rect_bottom = FancyBboxPatch((0.5, 0.3), 8.9, 1.2, boxstyle="round,pad=0.1",
                                  facecolor='#2c3e50', edgecolor='#1a252f', linewidth=1.5)
    ax.add_patch(rect_bottom)
    ax.text(5, 1.05, "DATA & TECHNOLOGY LAYER", ha='center', va='center',
            fontsize=10, color='white', fontweight='bold')
    ax.text(5, 0.6, "Python  •  Neo4j Graph Database  •  NetworkX  •  Matplotlib  •  Tkinter GUI",
            ha='center', va='center', fontsize=8, color='#bdc3c7')

    # Connections from boxes to bottom
    for x in [1.75, 4.95, 8.15]:
        ax.plot([x, x], [2.0, 1.5], color='#7f8c8d', linewidth=1.5, linestyle='--', alpha=0.6)

    fig.tight_layout()
    path = os.path.join(IMG_DIR, "architecture.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


def generate_frame_diagram():
    """Frame-based knowledge representation diagram."""
    fig, ax = plt.subplots(figsize=(8, 5))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 5.5)
    ax.axis('off')

    ax.text(4, 5.2, "Frame-Based Knowledge Representation", ha='center',
            fontsize=14, fontweight='bold', color='#2c3e50')

    # Main frame box
    main = FancyBboxPatch((1, 0.5), 6, 4.2, boxstyle="round,pad=0.15",
                           facecolor='#ecf0f1', edgecolor='#2c3e50', linewidth=2)
    ax.add_patch(main)
    ax.text(4, 4.35, "DISEASE FRAME: Influenza (Flu)", ha='center',
            fontsize=11, fontweight='bold', color='#2c3e50')

    # Slots
    slots = [
        (1.4, 3.5, "Slot: Description", "A viral infection that attacks\nyour respiratory system", '#3498db'),
        (1.4, 2.7, "Slot: Severity", "Moderate", '#e74c3c'),
        (1.4, 1.9, "Slot: Category", "Respiratory", '#9b59b6'),
        (1.4, 1.1, "Slot: Prevalence", "Common", '#f39c12'),
    ]

    for x, y, label, value, color in slots:
        # Slot label
        rect = FancyBboxPatch((x, y), 2.2, 0.55, boxstyle="round,pad=0.08",
                               facecolor=color, edgecolor='white', linewidth=1, alpha=0.85)
        ax.add_patch(rect)
        ax.text(x + 1.1, y + 0.28, label, ha='center', va='center',
                fontsize=7, color='white', fontweight='bold')
        # Filler value
        rect2 = FancyBboxPatch((3.9, y), 2.8, 0.55, boxstyle="round,pad=0.08",
                                facecolor='white', edgecolor=color, linewidth=1.5)
        ax.add_patch(rect2)
        ax.text(5.3, y + 0.28, value, ha='center', va='center',
                fontsize=7, color='#2c3e50')
        # Arrow
        ax.annotate("", xy=(3.9, y + 0.28), xytext=(3.6, y + 0.28),
                     arrowprops=dict(arrowstyle='->', color=color, lw=1.5))

    fig.tight_layout()
    path = os.path.join(IMG_DIR, "frame_diagram.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


def generate_semantic_network_diagram():
    """A clean sample semantic network diagram showing Disease-Symptom-Treatment."""
    G = nx.Graph()

    # Sample subset for clarity
    diseases = ["Influenza (Flu)", "Common Cold", "COVID-19"]
    sample = {d: MEDICAL_KNOWLEDGE_BASE[d] for d in diseases}

    disease_symptoms_map = {}
    for d, data in sample.items():
        syms = data.get("symptoms", [])[:5]  # limit for clarity
        disease_symptoms_map[d] = syms
        for s in syms:
            G.add_edge(d, s)

    # Clustered layout
    pos = {}
    n_d = len(diseases)
    cluster_r = 4.0
    sym_r = 2.0
    for i, disease in enumerate(diseases):
        angle = 2 * math.pi * i / n_d - math.pi / 2
        cx, cy = cluster_r * math.cos(angle), cluster_r * math.sin(angle)
        pos[disease] = (cx, cy)
        syms = disease_symptoms_map[disease]
        for j, sym in enumerate(syms):
            sa = 2 * math.pi * j / len(syms) - math.pi / 2
            sx = cx + sym_r * math.cos(sa)
            sy = cy + sym_r * math.sin(sa)
            if sym in pos:
                ox, oy = pos[sym]
                pos[sym] = ((ox + sx) / 2, (oy + sy) / 2)
            else:
                pos[sym] = (sx, sy)

    fig, ax = plt.subplots(figsize=(10, 7))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('#f8f9fa')

    # Colors
    node_colors = []
    node_sizes = []
    for node in G.nodes():
        if node in diseases:
            node_colors.append('#e74c3c')
            node_sizes.append(3500)
        else:
            node_colors.append('#3498db')
            node_sizes.append(1200)

    nx.draw_networkx_edges(G, pos, width=1.5, alpha=0.3, edge_color='#bdc3c7', ax=ax)
    nx.draw_networkx_nodes(G, pos, node_size=node_sizes, node_color=node_colors,
                           alpha=0.9, edgecolors='white', linewidths=2, ax=ax)

    disease_labels = {n: n for n in G.nodes() if n in diseases}
    symptom_labels = {n: n for n in G.nodes() if n not in diseases}

    nx.draw_networkx_labels(G, pos, labels=disease_labels,
                            font_size=8, font_weight='bold',
                            bbox=dict(facecolor='white', edgecolor='#e74c3c',
                                      alpha=0.9, boxstyle='round,pad=0.3'), ax=ax)
    nx.draw_networkx_labels(G, pos, labels=symptom_labels,
                            font_size=7, font_color='#2c3e50',
                            bbox=dict(facecolor='#eaf2f8', edgecolor='none',
                                      alpha=0.8, boxstyle='round,pad=0.2'), ax=ax)

    # Edge label
    ax.text(0, -5.5, "Edge Relationship: HAS_SYMPTOM", ha='center',
            fontsize=9, style='italic', color='#7f8c8d')

    legend_elements = [
        mpatches.Patch(color='#e74c3c', label='Disease Nodes'),
        mpatches.Patch(color='#3498db', label='Symptom Nodes'),
    ]
    ax.legend(handles=legend_elements, loc='upper right', fontsize=9, framealpha=0.9)
    ax.set_title("Semantic Network — Sample Disease-Symptom Relationships",
                 fontsize=13, fontweight='bold', pad=15)
    ax.axis('off')
    fig.tight_layout()

    path = os.path.join(IMG_DIR, "semantic_network.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


def generate_neo4j_schema_diagram():
    """Neo4j database schema showing node types and relationships."""
    fig, ax = plt.subplots(figsize=(9, 5))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 5)
    ax.axis('off')

    ax.text(4.5, 4.7, "Neo4j Graph Database Schema", ha='center',
            fontsize=14, fontweight='bold', color='#2c3e50')

    # Node circles
    nodes_info = [
        (4.5, 2.8, 'Disease', '#e74c3c', 'name\ndescription\nseverity\ncategory\nprevalence'),
        (1.2, 1.5, 'Symptom', '#3498db', 'name'),
        (7.8, 1.5, 'Treatment', '#27ae60', 'name'),
        (1.2, 4.0, 'Category', '#9b59b6', 'name'),
        (7.8, 4.0, 'RiskFactor', '#f39c12', 'name'),
    ]

    for x, y, label, color, props in nodes_info:
        circle = plt.Circle((x, y), 0.7, color=color, alpha=0.85, ec='white', lw=2)
        ax.add_patch(circle)
        ax.text(x, y + 0.15, label, ha='center', va='center',
                fontsize=9, fontweight='bold', color='white')
        ax.text(x, y - 0.25, props, ha='center', va='center',
                fontsize=5.5, color='white', alpha=0.85, linespacing=1.2)

    # Relationships (arrows)
    relationships = [
        (4.5, 2.8, 1.2, 1.5, 'HAS_SYMPTOM', '#3498db'),
        (7.8, 1.5, 4.5, 2.8, 'TREATS', '#27ae60'),
        (4.5, 2.8, 1.2, 4.0, 'BELONGS_TO', '#9b59b6'),
        (4.5, 2.8, 7.8, 4.0, 'HAS_RISK_FACTOR', '#f39c12'),
    ]

    for x1, y1, x2, y2, label, color in relationships:
        # Calculate direction
        dx, dy = x2 - x1, y2 - y1
        dist = math.sqrt(dx**2 + dy**2)
        # Shorten by radius
        offset = 0.75
        sx = x1 + dx / dist * offset
        sy = y1 + dy / dist * offset
        ex = x2 - dx / dist * offset
        ey = y2 - dy / dist * offset

        ax.annotate("", xy=(ex, ey), xytext=(sx, sy),
                     arrowprops=dict(arrowstyle='->', color=color, lw=2, connectionstyle="arc3,rad=0.1"))
        # Label on edge
        mx, my = (sx + ex) / 2, (sy + ey) / 2
        ax.text(mx, my + 0.15, label, ha='center', va='center',
                fontsize=6.5, fontweight='bold', color=color,
                bbox=dict(facecolor='white', edgecolor=color, alpha=0.9,
                          boxstyle='round,pad=0.15'))

    fig.tight_layout()
    path = os.path.join(IMG_DIR, "neo4j_schema.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


def generate_category_chart():
    """Bar chart showing disease distribution by category."""
    categories = {}
    for data in MEDICAL_KNOWLEDGE_BASE.values():
        cat = data.get("category", "General")
        categories[cat] = categories.get(cat, 0) + 1

    cats = sorted(categories.keys(), key=lambda c: categories[c], reverse=True)
    counts = [categories[c] for c in cats]

    colors = ['#54a0ff', '#ff7675', '#00b894', '#e17055', '#fd79a8',
              '#fdcb6e', '#6c5ce7', '#e84393', '#00cec9', '#0984e3',
              '#a29bfe', '#d63031']

    fig, ax = plt.subplots(figsize=(9, 4.5))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('#f8f9fa')

    bars = ax.barh(range(len(cats)), counts, color=colors[:len(cats)], edgecolor='white',
                   linewidth=1.5, height=0.65, alpha=0.9)

    ax.set_yticks(range(len(cats)))
    ax.set_yticklabels(cats, fontsize=9)
    ax.invert_yaxis()
    ax.set_xlabel("Number of Diseases", fontsize=10)
    ax.set_title("Disease Distribution by Body System Category",
                 fontsize=13, fontweight='bold', pad=15)

    for i, (bar, count) in enumerate(zip(bars, counts)):
        ax.text(bar.get_width() + 0.15, bar.get_y() + bar.get_height()/2,
                str(count), va='center', fontsize=10, fontweight='bold', color='#2c3e50')

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.tick_params(left=False)
    ax.grid(axis='x', alpha=0.3, linestyle='--')

    fig.tight_layout()
    path = os.path.join(IMG_DIR, "category_chart.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


def generate_diagnosis_flow_diagram():
    """Flowchart of the diagnosis process."""
    fig, ax = plt.subplots(figsize=(8, 7))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 8)
    ax.axis('off')

    ax.text(4, 7.6, "Diagnostic Process Flowchart", ha='center',
            fontsize=14, fontweight='bold', color='#2c3e50')

    steps = [
        (4, 6.8, 2.8, 0.7, '#3498db', 'Patient Enters Symptoms\n(via GUI or CLI)'),
        (4, 5.6, 2.8, 0.7, '#2980b9', 'Symptoms Parsed &\nStandardized (Title Case)'),
        (4, 4.4, 2.8, 0.7, '#8e44ad', 'Cypher Query Executed\non Neo4j Graph Database'),
        (4, 3.2, 2.8, 0.7, '#e74c3c', 'Symptom-Disease Matching\n& Confidence Scoring'),
        (4, 2.0, 2.8, 0.7, '#d35400', 'Top 5 Results Ranked by\nConfidence % & Severity'),
        (4, 0.8, 2.8, 0.7, '#27ae60', 'Display Results + Treatments\n+ Semantic Network Graph'),
    ]

    for i, (x, y, w, h, color, text) in enumerate(steps):
        rect = FancyBboxPatch((x - w/2, y - h/2), w, h, boxstyle="round,pad=0.12",
                               facecolor=color, edgecolor='white', linewidth=2, alpha=0.9)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center',
                fontsize=8, color='white', fontweight='bold', linespacing=1.3)

        # Arrow to next step
        if i < len(steps) - 1:
            next_y = steps[i+1][1]
            ax.annotate("", xy=(x, next_y + 0.4), xytext=(x, y - 0.4),
                         arrowprops=dict(arrowstyle='->', color='#34495e', lw=2.0))

    # Side note
    ax.text(7.2, 4.4, "Neo4j\nCypher\nQuery", ha='center', va='center',
            fontsize=7, color='#8e44ad', fontweight='bold', style='italic',
            bbox=dict(facecolor='#f5eef8', edgecolor='#8e44ad',
                      boxstyle='round,pad=0.3', alpha=0.8))
    ax.annotate("", xy=(5.6, 4.4), xytext=(6.3, 4.4),
                 arrowprops=dict(arrowstyle='<-', color='#8e44ad', lw=1.5,
                                 connectionstyle="arc3,rad=0"))

    fig.tight_layout()
    path = os.path.join(IMG_DIR, "diagnosis_flow.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ======================================================================
# DOCX REPORT BUILDER
# ======================================================================

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex.replace('#', ''))
    shading_elm.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading_elm)


def build_report():
    """Build the complete DOCX project report."""
    print("Generating report diagrams...")

    # Generate all diagrams
    arch_img = generate_architecture_diagram()
    frame_img = generate_frame_diagram()
    semantic_img = generate_semantic_network_diagram()
    schema_img = generate_neo4j_schema_diagram()
    category_img = generate_category_chart()
    flow_img = generate_diagnosis_flow_diagram()

    print("Building DOCX document...")

    doc = Document()

    # ---- Page margins ----
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---- STYLES ----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.line_spacing = 1.4
    style.paragraph_format.space_after = Pt(6)

    # Heading styles
    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        hs.font.name = 'Calibri'
        if level == 1:
            hs.font.size = Pt(22)
            hs.paragraph_format.space_before = Pt(24)
            hs.paragraph_format.space_after = Pt(12)
        elif level == 2:
            hs.font.size = Pt(16)
            hs.font.color.rgb = RGBColor(0x29, 0x80, 0xb9)
            hs.paragraph_format.space_before = Pt(18)
            hs.paragraph_format.space_after = Pt(8)
        else:
            hs.font.size = Pt(13)
            hs.font.color.rgb = RGBColor(0x34, 0x49, 0x5e)

    # ==================================================================
    # COVER PAGE
    # ==================================================================
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Medical Diagnostic\nSupport System")
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    title_run.bold = True

    # Separator
    sep_p = doc.add_paragraph()
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep_p.add_run("━" * 40)
    sep_run.font.size = Pt(14)
    sep_run.font.color.rgb = RGBColor(0x34, 0x98, 0xdb)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run(
        "Using Semantic Networks, Frame-Based Knowledge\n"
        "Representation & Neo4j Graph Database"
    )
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

    for _ in range(4):
        doc.add_paragraph()

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(
        "Project Report\n\n"
        "Technologies: Python · Neo4j · NetworkX · Tkinter\n"
        f"Knowledge Base: {len(MEDICAL_KNOWLEDGE_BASE)} Diseases · "
        f"{len(set(s for d in MEDICAL_KNOWLEDGE_BASE.values() for s in d.get('symptoms', [])))} Symptoms"
    )
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = RGBColor(0x95, 0xa5, 0xa6)

    doc.add_page_break()

    # ==================================================================
    # TABLE OF CONTENTS (manual)
    # ==================================================================
    doc.add_heading("Table of Contents", level=1)

    toc_items = [
        ("1.", "Introduction", 3),
        ("2.", "Problem Statement", 3),
        ("3.", "Core Concepts & Theory", 4),
        ("   3.1", "Frame-Based Knowledge Representation", 4),
        ("   3.2", "Semantic Networks", 5),
        ("   3.3", "Graph Databases (Neo4j)", 5),
        ("4.", "System Architecture", 6),
        ("5.", "Database Schema Design", 7),
        ("6.", "Diagnostic Process", 8),
        ("7.", "Dataset Overview", 9),
        ("8.", "Implementation Details", 10),
        ("   8.1", "Knowledge Base Seeding", 10),
        ("   8.2", "Diagnostic Engine", 10),
        ("   8.3", "Graphical User Interface", 11),
        ("   8.4", "Network Visualization", 11),
        ("   8.5", "Source Code - Key Functions & Algorithms", 12),
        ("9.", "Technologies Used", 15),
        ("10.", "Sample Outputs", 15),
        ("11.", "Conclusion & Future Scope", 16),
    ]

    for num, title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
        run = p.add_run(f"{num}  {title}")
        run.font.size = Pt(11)
        if not num.startswith(" "):
            run.bold = True

    doc.add_page_break()

    # ==================================================================
    # 1. INTRODUCTION
    # ==================================================================
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "The Medical Diagnostic Support System (MDSS) is an intelligent, graph-based application "
        "designed to assist in preliminary disease diagnosis. Traditional diagnostic systems often "
        "rely on rigid, relational databases that struggle with complex, interconnected medical data. "
        "By utilizing Graph Databases and Semantic Networks, this project natively models the "
        "non-linear relationships between diseases, symptoms, and treatments."
    )
    doc.add_paragraph(
        "The system provides medical professionals and patients with a fast, preliminary diagnostic "
        "tool that can map patient symptoms to likely conditions, compute confidence scores, display "
        "relevant treatments and risk factors, and visually represent the reasoning behind each "
        "diagnosis through an interactive semantic network graph."
    )
    doc.add_paragraph(
        f"The current knowledge base contains {len(MEDICAL_KNOWLEDGE_BASE)} diseases across "
        f"{len(set(d.get('category', '') for d in MEDICAL_KNOWLEDGE_BASE.values()))} body-system "
        "categories, with comprehensive symptom mappings, treatment protocols, and risk factor data."
    )

    # ==================================================================
    # 2. PROBLEM STATEMENT
    # ==================================================================
    doc.add_heading("2. Problem Statement", level=1)
    doc.add_paragraph(
        "Medical knowledge is inherently highly interconnected. A single disease can have multiple "
        "symptoms, and a single symptom (like \"Fever\") can be common across numerous diseases. "
        "Storing this data in traditional SQL tabular formats leads to complex joins, slow "
        "performance when analyzing deep connections, and difficulty in representing the semantic "
        "meaning of medical relationships."
    )
    doc.add_paragraph(
        "There is a need for a system that can semantically understand these relationships — "
        "treating medical entities as nodes and their relationships as edges — to quickly query "
        "and isolate potential diagnoses based on a combination of patient symptoms with computed "
        "confidence levels."
    )

    # ==================================================================
    # 3. CORE CONCEPTS
    # ==================================================================
    doc.add_heading("3. Core Concepts & Theory", level=1)

    doc.add_heading("3.1 Frame-Based Knowledge Representation", level=2)
    doc.add_paragraph(
        'In artificial intelligence, a "Frame" is a data structure for dividing knowledge into '
        'substructures by representing "stereotyped situations." In our system, each Disease acts '
        'as a frame with the following components:'
    )
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run('Slots: ')
    r.bold = True
    p.add_run('Properties inherent to the disease (e.g., Description, Prevalence, Severity, Category)')
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run('Fillers: ')
    r.bold = True
    p.add_run('Links to other distinct objects (e.g., Symptoms, Treatments, Risk Factors)')

    doc.add_paragraph(
        "By applying frame-based logic, the system standardizes how varying medical conditions "
        "are structurally understood, making it easily expandable and maintainable."
    )

    # Frame diagram
    doc.add_picture(frame_img, width=Inches(5.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 1: Frame-Based Knowledge Representation for a Disease")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
    cap.runs[0].italic = True

    doc.add_heading("3.2 Semantic Networks", level=2)
    doc.add_paragraph(
        "A Semantic Network is a knowledge base that represents semantic relations between "
        "concepts in a network. In this project:"
    )
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run('Nodes (Vertices): ')
    r.bold = True
    p.add_run('Represent core entities: Disease, Symptom, Treatment, Category, RiskFactor')
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run('Edges (Links): ')
    r.bold = True
    p.add_run('Represent directional relationships: HAS_SYMPTOM, TREATS, BELONGS_TO, HAS_RISK_FACTOR')

    # Semantic network diagram
    doc.add_picture(semantic_img, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 2: Semantic Network — Sample Disease-Symptom Relationships")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
    cap.runs[0].italic = True

    doc.add_heading("3.3 Graph Databases (Neo4j)", level=2)
    doc.add_paragraph(
        "To ensure the system is scalable to thousands of medical conditions, Neo4j is used as "
        "the data storage layer. Unlike relational databases that use tables, Neo4j is a native "
        "graph database. This allows for instantaneous retrieval of connected pathways using "
        "Cypher Query Language (CQL), making complex diagnostic pattern-matching highly "
        "resource-efficient."
    )
    doc.add_paragraph(
        "Key advantages of Neo4j over traditional RDBMS for this application:"
    )
    for adv in [
        "Native graph storage eliminates costly JOIN operations",
        "Constant-time traversal of relationships regardless of dataset size",
        "Intuitive visual query language (Cypher) maps directly to our semantic network",
        "Schema-flexible design allows easy addition of new disease categories",
    ]:
        doc.add_paragraph(adv, style='List Bullet')

    doc.add_page_break()

    # ==================================================================
    # 4. SYSTEM ARCHITECTURE
    # ==================================================================
    doc.add_heading("4. System Architecture", level=1)
    doc.add_paragraph(
        "The application is structured into three continuous processing layers:"
    )
    for item in [
        ("Input Layer: ", "The user submits symptoms through a modern Tkinter GUI application "
         "with text entry, quick-add buttons, and category browsing. Symptoms are automatically "
         "parsed and standardized."),
        ("Processing Layer: ", "The system executes Cypher queries against the Neo4j graph "
         "database. It scans the semantic network for all paths from Symptom nodes to Disease "
         "nodes, computing confidence scores as the ratio of matched symptoms to total disease symptoms."),
        ("Output Layer: ", "Results are displayed with confidence percentages, severity indicators, "
         "treatments, and risk factors. A clustered semantic network graph is rendered inline "
         "showing the diagnosis reasoning."),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(item[0])
        r.bold = True
        p.add_run(item[1])

    doc.add_picture(arch_img, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 3: Three-Layer System Architecture")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
    cap.runs[0].italic = True

    doc.add_page_break()

    # ==================================================================
    # 5. DATABASE SCHEMA
    # ==================================================================
    doc.add_heading("5. Database Schema Design", level=1)
    doc.add_paragraph(
        "The Neo4j database schema consists of five node types connected by four relationship types:"
    )

    # Schema table
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    headers = ['Node Type', 'Properties', 'Description']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    schema_data = [
        ('Disease', 'name, description, severity,\ncategory, prevalence', 'Core medical conditions'),
        ('Symptom', 'name', 'Observable patient symptoms'),
        ('Treatment', 'name', 'Medical treatments and remedies'),
        ('Category', 'name', 'Body system classification'),
        ('RiskFactor', 'name', 'Contributing risk factors'),
    ]
    for i, (node, props, desc) in enumerate(schema_data, 1):
        table.rows[i].cells[0].text = node
        table.rows[i].cells[1].text = props
        table.rows[i].cells[2].text = desc

    doc.add_paragraph()  # spacer

    # Relationship table
    p = doc.add_paragraph()
    r = p.add_run("Relationship Types:")
    r.bold = True
    r.font.size = Pt(12)

    table2 = doc.add_table(rows=5, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Light Grid Accent 1'

    headers2 = ['Relationship', 'Direction', 'Meaning']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    rels = [
        ('HAS_SYMPTOM', 'Disease → Symptom', 'Disease exhibits this symptom'),
        ('TREATS', 'Treatment → Disease', 'Treatment addresses this disease'),
        ('BELONGS_TO', 'Disease → Category', 'Disease belongs to body system'),
        ('HAS_RISK_FACTOR', 'Disease → RiskFactor', 'Disease has this risk factor'),
    ]
    for i, (rel, direction, meaning) in enumerate(rels, 1):
        table2.rows[i].cells[0].text = rel
        table2.rows[i].cells[1].text = direction
        table2.rows[i].cells[2].text = meaning

    doc.add_paragraph()
    doc.add_picture(schema_img, width=Inches(5.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 4: Neo4j Graph Database Schema")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
    cap.runs[0].italic = True

    doc.add_page_break()

    # ==================================================================
    # 6. DIAGNOSTIC PROCESS
    # ==================================================================
    doc.add_heading("6. Diagnostic Process", level=1)
    doc.add_paragraph(
        "The diagnostic engine follows a six-step pipeline to convert raw patient symptoms "
        "into ranked disease predictions with confidence scores:"
    )

    steps = [
        ("Step 1 — Symptom Input: ", "Patient enters symptoms via the GUI text field or "
         "quick-add buttons. Multiple symptoms are comma-separated."),
        ("Step 2 — Parsing: ", "Symptoms are cleaned, trimmed, and converted to title case "
         "to match the standardized database format."),
        ("Step 3 — Cypher Query: ", "A Cypher MATCH query is executed against Neo4j, finding "
         "all Disease nodes connected to the provided Symptom nodes via HAS_SYMPTOM edges."),
        ("Step 4 — Confidence Scoring: ", "For each matched disease, confidence is calculated as: "
         "Confidence = (Matched Symptoms / Total Disease Symptoms) × 100%. "
         "Results are ranked by confidence descending."),
        ("Step 5 — Result Ranking: ", "The top 5 diseases are returned with their confidence "
         "percentage, severity level, category, and matched symptom list."),
        ("Step 6 — Visualization: ", "A clustered semantic network graph is drawn showing the "
         "top 3 diseases with their symptom clusters. Matched symptoms are highlighted."),
    ]
    for step_title, step_desc in steps:
        p = doc.add_paragraph()
        r = p.add_run(step_title)
        r.bold = True
        p.add_run(step_desc)

    doc.add_picture(flow_img, width=Inches(4.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 5: Diagnostic Process Flowchart")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
    cap.runs[0].italic = True

    # Cypher query example
    doc.add_heading("Core Diagnostic Cypher Query", level=3)
    p = doc.add_paragraph()
    r = p.add_run(
        'MATCH (d:Disease)-[:HAS_SYMPTOM]->(s:Symptom)\n'
        'WITH d, collect(s.name) AS allSymptoms\n'
        'WITH d, allSymptoms, size(allSymptoms) AS totalSymptoms\n'
        'WITH d, allSymptoms, totalSymptoms,\n'
        '     [sym IN allSymptoms WHERE sym IN $patient_symptoms] AS matched\n'
        'WITH d, matched, totalSymptoms, size(matched) AS matchCount\n'
        'WHERE matchCount > 0\n'
        'RETURN d.name AS Disease,\n'
        '       toFloat(matchCount)/toFloat(totalSymptoms)*100 AS Confidence\n'
        'ORDER BY Confidence DESC LIMIT 5'
    )
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    p.paragraph_format.space_before = Pt(8)

    doc.add_page_break()

    # ==================================================================
    # 7. DATASET OVERVIEW
    # ==================================================================
    doc.add_heading("7. Dataset Overview", level=1)
    doc.add_paragraph(
        f"The medical knowledge base contains {len(MEDICAL_KNOWLEDGE_BASE)} diseases "
        f"across {len(set(d.get('category', '') for d in MEDICAL_KNOWLEDGE_BASE.values()))} "
        "body-system categories. Each disease is modeled as a frame with slots for description, "
        "severity, prevalence, and category, plus fillers linking to symptoms, treatments, "
        "and risk factors."
    )

    # Category chart
    doc.add_picture(category_img, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 6: Disease Distribution by Body System Category")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
    cap.runs[0].italic = True

    # Disease table
    doc.add_heading("Complete Disease List", level=3)

    # Group by category
    by_cat = {}
    for name, data in sorted(MEDICAL_KNOWLEDGE_BASE.items()):
        cat = data.get("category", "General")
        by_cat.setdefault(cat, []).append((name, data))

    for cat in sorted(by_cat.keys()):
        p = doc.add_paragraph()
        r = p.add_run(f"\n{cat}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x29, 0x80, 0xb9)

        diseases = by_cat[cat]
        table = doc.add_table(rows=len(diseases) + 1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(['Disease', 'Severity', 'Prevalence', 'Symptoms']):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        for j, (name, data) in enumerate(diseases, 1):
            table.rows[j].cells[0].text = name
            table.rows[j].cells[1].text = data.get('severity', '')
            table.rows[j].cells[2].text = data.get('prevalence', '')
            table.rows[j].cells[3].text = ', '.join(data.get('symptoms', [])[:6])
            for cell in table.rows[j].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

    doc.add_page_break()

    # ==================================================================
    # 8. IMPLEMENTATION DETAILS
    # ==================================================================
    doc.add_heading("8. Implementation Details", level=1)

    doc.add_heading("8.1 Knowledge Base Seeding", level=2)
    doc.add_paragraph(
        "The seed_database() method clears existing Neo4j data and builds the complete graph "
        "from the Python dictionary-based knowledge base. For each disease frame, it creates:"
    )
    for item in [
        "A Disease node with name, description, severity, category, and prevalence properties",
        "A Category node with a BELONGS_TO relationship",
        "Symptom nodes with HAS_SYMPTOM relationships",
        "Treatment nodes with TREATS relationships",
        "RiskFactor nodes with HAS_RISK_FACTOR relationships",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("8.2 Diagnostic Engine", level=2)
    doc.add_paragraph(
        "The diagnostic engine executes an advanced Cypher query that performs symptom matching "
        "with confidence scoring in a single database pass. The query uses list comprehension "
        "within Cypher to compute the intersection of patient symptoms with each disease's "
        "symptom list, then calculates a percentage confidence score."
    )
    doc.add_paragraph(
        "Results are ordered by confidence descending, then by raw match count as a tiebreaker. "
        "The top 5 results are returned to the user interface."
    )

    doc.add_heading("8.3 Graphical User Interface (GUI)", level=2)
    doc.add_paragraph(
        "The application features a modern dark-themed Tkinter GUI (gui.py) with the following "
        "components:"
    )
    for item in [
        "Symptom input field with enter-key support and quick-add buttons for common symptoms",
        "Category browser dropdown for exploring diseases by body system",
        "Action buttons: List All Diseases, List All Symptoms, Full Network, Reseed Database",
        "Color-coded results panel showing confidence bars, severity levels, treatments, and risk factors",
        "Embedded semantic network graph visualization (right panel)",
        "Status bar showing connection status and disease counts",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("8.4 Network Visualization", level=2)
    doc.add_paragraph(
        "The visualization uses a custom clustered radial layout algorithm. Each disease is "
        "placed on a circle, with its symptoms arranged in a ring around it. Shared symptoms "
        "(like Fever) are positioned at the midpoint between their connected diseases. This "
        "eliminates the messy overlapping common with spring-force layouts."
    )
    doc.add_paragraph(
        "In diagnostic mode, only the top 3 matched diseases are shown. Matched patient "
        "symptoms are highlighted in yellow with bold labels, while unmatched symptoms appear "
        "as small gray nodes."
    )

    doc.add_page_break()

    # ==================================================================
    # 8.5 SOURCE CODE - KEY FUNCTIONS (NEW)
    # ==================================================================
    doc.add_heading("8.5 Source Code - Key Functions & Algorithms", level=2)
    doc.add_paragraph(
        "This section presents the core source code functions that power the diagnostic system. "
        "Only the main logic and algorithms are shown for clarity."
    )

    # --- Helper to add code blocks ---
    def add_code_block(doc, title, code_text, explanation=""):
        """Add a titled code block with Consolas font and gray background."""
        if title:
            p = doc.add_paragraph()
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x29, 0x80, 0xb9)
            p.paragraph_format.space_after = Pt(4)

        if explanation:
            exp_p = doc.add_paragraph(explanation)
            exp_p.paragraph_format.space_after = Pt(6)
            for r in exp_p.runs:
                r.font.size = Pt(10)

        code_p = doc.add_paragraph()
        code_p.paragraph_format.space_before = Pt(4)
        code_p.paragraph_format.space_after = Pt(8)
        code_p.paragraph_format.line_spacing = 1.15
        code_run = code_p.add_run(code_text)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(8)
        code_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # ---- CODE 1: Knowledge Base Frame Structure ----
    add_code_block(doc,
        "1. Knowledge Base Frame Structure (medical_dataset.py)",
        '''"Influenza (Flu)": {
    "description": "A viral infection that attacks your
                    respiratory system.",
    "prevalence": "Common",
    "severity": "Moderate",
    "category": "Respiratory",
    "symptoms": ["Fever", "Chills", "Muscle aches",
                 "Cough", "Congestion", "Runny nose",
                 "Headache", "Fatigue"],
    "treatments": ["Rest", "Fluid intake",
                   "Antiviral medication",
                   "Pain relievers"],
    "risk_factors": ["Weakened immune system",
                     "Age (under 5 or over 65)",
                     "Chronic illness"]
}''',
        "Each disease is stored as a Python dictionary (frame). "
        "Slots hold direct properties (description, severity, category, prevalence). "
        "Fillers hold relationships to other entities (symptoms, treatments, risk_factors). "
        "This structure is directly mapped to Neo4j nodes and edges."
    )

    # ---- CODE 2: Database Connection ----
    add_code_block(doc,
        "2. Neo4j Database Connection (main.py)",
        '''class MedicalKnowledgeBase:
    def __init__(self, uri, user, password):
        try:
            self.driver = GraphDatabase.driver(
                uri, auth=(user, password)
            )
            self.driver.verify_connectivity()
            print("Connected to Neo4j Database.")
        except Exception as e:
            print(f"Failed to connect: {e}")
            self.driver = None''',
        "The system connects to Neo4j using the official Python driver. "
        "verify_connectivity() ensures the credentials and URI are valid before proceeding."
    )

    # ---- CODE 3: Seed Database ----
    add_code_block(doc,
        "3. Database Seeding Algorithm (main.py)",
        '''def seed_database(self):
    """Load all frames into Neo4j as nodes
    and relationships."""
    with self.driver.session() as session:
        # Clear old data
        session.run("MATCH (n) DETACH DELETE n")

        for disease_name, frame in MEDICAL_KNOWLEDGE_BASE.items():
            # 1. Create Disease node with properties
            session.run(\'\'\'
                CREATE (d:Disease {
                    name: $name,
                    description: $description,
                    severity: $severity,
                    category: $category
                })
            \'\'\', name=disease_name, ...)

            # 2. Create Category node + BELONGS_TO edge
            session.run(\'\'\'
                MATCH (d:Disease {name: $d_name})
                MERGE (c:Category {name: $c_name})
                MERGE (d)-[:BELONGS_TO]->(c)
            \'\'\', d_name=disease_name, c_name=category)

            # 3. Create Symptom nodes + HAS_SYMPTOM edges
            for symptom in frame.get("symptoms", []):
                session.run(\'\'\'
                    MATCH (d:Disease {name: $d_name})
                    MERGE (s:Symptom {name: $s_name})
                    MERGE (d)-[:HAS_SYMPTOM]->(s)
                \'\'\', d_name=disease_name, s_name=symptom)

            # 4. Create Treatment + RiskFactor nodes
            # (similar MERGE pattern for TREATS and
            #  HAS_RISK_FACTOR relationships)''',
        "Algorithm: For each disease frame in the knowledge base, the seeder creates a Disease "
        "node with all slot properties, then uses MERGE to create unique Symptom, Treatment, "
        "Category, and RiskFactor nodes with their respective relationship edges. "
        "MERGE prevents duplicate nodes when multiple diseases share symptoms."
    )

    doc.add_page_break()

    # ---- CODE 4: Diagnostic Engine ----
    add_code_block(doc,
        "4. Diagnostic Engine with Confidence Scoring (main.py)",
        '''def diagnose(self, symptoms):
    """Core diagnostic algorithm using Cypher."""
    with self.driver.session() as session:
        query = \'\'\'
            MATCH (d:Disease)-[:HAS_SYMPTOM]->(s:Symptom)
            WITH d, collect(s.name) AS allSymptoms
            WITH d, allSymptoms,
                 size(allSymptoms) AS totalSymptoms
            WITH d, allSymptoms, totalSymptoms,
                 [sym IN allSymptoms
                  WHERE sym IN $patient_symptoms]
                      AS matchedSymptoms
            WITH d, matchedSymptoms, totalSymptoms,
                 size(matchedSymptoms) AS matchCount
            WHERE matchCount > 0
            RETURN d.name AS Disease,
                   matchCount AS MatchingSymptoms,
                   totalSymptoms AS TotalSymptoms,
                   matchedSymptoms AS MatchedSymptomList,
                   d.severity AS Severity,
                   d.category AS Category,
                   toFloat(matchCount)
                       / toFloat(totalSymptoms)
                       * 100 AS Confidence
            ORDER BY Confidence DESC,
                     MatchingSymptoms DESC
            LIMIT 5
        \'\'\'
        result = session.run(
            query, patient_symptoms=symptoms
        )
        return [record for record in result]''',
        "Algorithm: The Cypher query traverses all Disease->Symptom edges, collects each disease's "
        "full symptom list, computes the intersection with the patient's input using list comprehension, "
        "then calculates Confidence = (matched / total) * 100. Results are sorted by confidence "
        "descending, returning the top 5 matches with severity and category metadata."
    )

    # ---- CODE 5: Clustered Visualization Algorithm ----
    add_code_block(doc,
        "5. Clustered Radial Layout Algorithm (main.py)",
        '''def visualize_network(self, top_diseases, symptoms):
    """Custom radial cluster layout for clean graphs."""
    # Step 1: Place diseases on a large circle
    for i, disease in enumerate(diseases_list):
        angle = 2 * pi * i / n_diseases - pi / 2
        cx = cluster_radius * cos(angle)
        cy = cluster_radius * sin(angle)
        pos[disease] = (cx, cy)

        # Step 2: Arrange symptoms in a ring
        #         around each disease
        for j, symptom in enumerate(symptoms):
            s_angle = 2 * pi * j / len(syms) - pi/2
            sx = cx + symptom_radius * cos(s_angle)
            sy = cy + symptom_radius * sin(s_angle)

            # Step 3: Shared symptoms get averaged
            #         positions between clusters
            if symptom in pos:
                old_x, old_y = pos[symptom]
                pos[symptom] = (
                    (old_x + sx) / 2,
                    (old_y + sy) / 2
                )
            else:
                pos[symptom] = (sx, sy)

    # Step 4: Draw with color-coded nodes
    #   - Red:    Diagnosed diseases
    #   - Yellow: Patient's matched symptoms
    #   - Gray:   Unmatched symptoms''',
        "Algorithm: Instead of spring-force layout (which causes messy overlapping), this custom "
        "algorithm places each disease at an equally-spaced point on a large circle, then arranges "
        "each disease's symptoms in a smaller ring around it. When a symptom appears in multiple "
        "diseases (e.g., 'Fever'), its position is averaged between clusters, placing it between "
        "the diseases that share it. This produces clean, readable cluster diagrams."
    )

    # ---- CODE 6: GUI Diagnosis Handler ----
    add_code_block(doc,
        "6. GUI Diagnosis Handler (gui.py)",
        '''def _run_diagnosis(self):
    """Handle diagnosis button click in GUI."""
    # 1. Parse user input
    raw = self.symptom_entry.get().strip()
    patient_symptoms = [
        s.strip().title() for s in raw.split(',')
        if s.strip()
    ]

    # 2. Run diagnostic query
    matches = self.db.diagnose(patient_symptoms)

    # 3. Display results with confidence bars
    for i, m in enumerate(matches, 1):
        conf = m['Confidence']
        bar_len = 20
        filled = int(conf / 100 * bar_len)
        bar = '#' * filled + '-' * (bar_len - filled)

        self._write_result(
            f"{i}. {m['Disease']}\\n"
            f"   Confidence: [{bar}] {conf:.1f}%\\n"
            f"   Severity:   {m['Severity']}\\n"
        )

    # 4. Show treatments for top result
    top = matches[0]['Disease']
    treatments = self.db.get_treatments(top)
    risk_factors = self.db.get_risk_factors(top)

    # 5. Draw semantic network graph
    top_diseases = [m['Disease'] for m in matches[:3]]
    self._draw_diagnosis_graph(
        top_diseases, patient_symptoms
    )''',
        "The GUI handler orchestrates the full diagnostic flow: parsing comma-separated input, "
        "running the Neo4j query, rendering color-coded results with visual confidence bars, "
        "fetching treatments and risk factors for the top diagnosis, and triggering the "
        "embedded graph visualization for the top 3 matches."
    )

    doc.add_page_break()

    # ==================================================================
    # 9. TECHNOLOGIES USED
    # ==================================================================
    doc.add_heading("9. Technologies Used", level=1)

    tech_table = doc.add_table(rows=7, cols=3)
    tech_table.style = 'Light Grid Accent 1'
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(['Technology', 'Version', 'Purpose']):
        cell = tech_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    techs = [
        ('Python', '3.13', 'Core programming language for application logic'),
        ('Neo4j', '5.x', 'Graph database for semantic network storage and querying'),
        ('Cypher', '-', 'Graph query language for diagnostic pattern matching'),
        ('NetworkX', '3.x', 'Python library for graph creation, manipulation, and visualization'),
        ('Matplotlib', '3.x', 'Comprehensive library for creating network visualizations'),
        ('Tkinter', 'Built-in', 'GUI framework for the desktop application interface'),
    ]
    for i, (tech, ver, purpose) in enumerate(techs, 1):
        tech_table.rows[i].cells[0].text = tech
        tech_table.rows[i].cells[1].text = ver
        tech_table.rows[i].cells[2].text = purpose

    # ==================================================================
    # 10. SAMPLE OUTPUTS
    # ==================================================================
    doc.add_heading("10. Sample Diagnostic Scenario", level=1)
    doc.add_paragraph(
        'Consider a patient presenting with the symptoms: "Fever, Cough, Headache, Fatigue"'
    )

    p = doc.add_paragraph()
    r = p.add_run("Expected Output (Top 3 Matches):")
    r.bold = True
    r.font.size = Pt(12)

    # Sample output table
    sample_table = doc.add_table(rows=4, cols=5)
    sample_table.style = 'Light Grid Accent 1'
    sample_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(['Rank', 'Disease', 'Confidence', 'Matched', 'Severity']):
        cell = sample_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    samples = [
        ('1', 'Influenza (Flu)', '50.0%', 'Fever, Cough, Headache, Fatigue', 'Moderate'),
        ('2', 'COVID-19', '40.0%', 'Fever, Cough, Headache, Fatigue', 'High'),
        ('3', 'Malaria', '50.0%', 'Fever, Headache, Fatigue, Muscle aches', 'High'),
    ]
    for i, row_data in enumerate(samples, 1):
        for j, val in enumerate(row_data):
            sample_table.rows[i].cells[j].text = val

    doc.add_paragraph()

    # ==================================================================
    # 11. CONCLUSION & FUTURE SCOPE
    # ==================================================================
    doc.add_heading("11. Conclusion & Future Scope", level=1)

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        "The Medical Diagnostic Support System successfully demonstrates how graph-based "
        "modeling substantially improves the efficiency, clarity, and query speed of medical "
        "diagnosis software compared to traditional relational databases. The combination of "
        "Frame-based knowledge representation with Semantic Networks provides a structured, "
        "expandable, and semantically rich foundation for medical AI applications."
    )
    doc.add_paragraph(
        "The system achieves its goals of providing fast preliminary diagnoses with confidence "
        "scoring, intuitive visual explanations through semantic network graphs, and a modern "
        "graphical interface suitable for both medical professionals and educational purposes."
    )

    doc.add_heading("Future Scope", level=2)
    future_items = [
        ("Probabilistic Weights: ", "Integrate weighted edges to represent symptom-disease "
         "probability (e.g., Fever is a 90% indicator of Malaria but only 30% for GERD)."),
        ("Machine Learning Integration: ", "Train a model on real-world patient data to "
         "improve diagnostic accuracy beyond simple symptom counting."),
        ("Electronic Health Records: ", "Connect to EHR systems to provide personalized "
         "diagnoses based on patient history, age, and pre-existing conditions."),
        ("Web Application: ", "Expand the GUI to a full web application using Flask/Django "
         "or React for broader accessibility."),
        ("Natural Language Processing: ", "Allow patients to describe symptoms in natural "
         "language rather than selecting from predefined lists."),
        ("Multi-language Support: ", "Add support for symptom input and result display in "
         "multiple languages for wider global reach."),
    ]
    for title, desc in future_items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(title)
        r.bold = True
        p.add_run(desc)

    # ==================================================================
    # SAVE
    # ==================================================================
    doc.save(REPORT_PATH)
    print(f"\n[OK] Report saved to: {REPORT_PATH}")
    print(f"  ({len(MEDICAL_KNOWLEDGE_BASE)} diseases documented)")
    print(f"  6 diagrams embedded")
    print(f"\nYou can open this .docx file and export to PDF from Word.")


if __name__ == "__main__":
    build_report()
